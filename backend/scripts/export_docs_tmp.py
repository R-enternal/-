"""临时脚本：提取两份策划文档文本（UTF-8 输出到临时文件）"""

import docx
import pdfplumber

DOCX = r"D:\电脑管家迁移文件\微信聊天记录搬家\WeChat Files\wxid_xlfaftb3juwb12\FileStorage\Temp\Copy\仓维云商业计划书(1).docx"
PDF = r"D:\电脑管家迁移文件\微信聊天记录搬家\WeChat Files\wxid_xlfaftb3juwb12\FileStorage\Temp\Copy\策划书（pdf版）.pdf"
OUT1 = r"D:\AI Program\仓维云-agent\tmp_plan1.txt"
OUT2 = r"D:\AI Program\仓维云-agent\tmp_plan2.txt"


def main() -> None:
    # docx
    doc = docx.Document(DOCX)
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for i, table in enumerate(doc.tables):
        out.append(f"[表格{i + 1}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            out.append(" | ".join(cells))
    with open(OUT1, "w", encoding="utf-8") as f:
        f.write("\n".join(out))
    print(f"docx 导出完成，共 {len(out)} 段")

    # pdf
    out2 = []
    with pdfplumber.open(PDF) as pdf:
        print("PDF 总页数:", len(pdf.pages))
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            out2.append(f"\n===== 第{i + 1}页 =====\n{text}")
    with open(OUT2, "w", encoding="utf-8") as f:
        f.write("\n".join(out2))
    print("pdf 导出完成")


if __name__ == "__main__":
    main()
