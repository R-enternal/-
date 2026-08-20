"""知识库服务：多格式文档解析 → 结构感知切块 → 向量化（GLM embedding）→ 检索（向量 + BM25 混合）"""

import csv
import io
import json
import math
import re
from collections import Counter
from pathlib import Path
from typing import Any, Callable

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger
from pydantic import SecretStr
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import config
from app.database import SessionLocal
from app.models.kb import KbDocument

_COLLECTION = "kb_docs"
_vector_store: Chroma | None = None
_corpus_cache: list[dict[str, Any]] | None = None  # 混合检索用的全文片段缓存


# ---------------------------------------------------------------------------
# 向量库
# ---------------------------------------------------------------------------
def _get_embeddings() -> OpenAIEmbeddings:
    """智谱 GLM embedding（OpenAI 兼容）"""
    return OpenAIEmbeddings(
        model=config.embedding_model,
        api_key=SecretStr(config.embedding_api_key),
        base_url=config.embedding_base_url,
    )


def _get_vector_store() -> Chroma:
    """Chroma 持久化向量库（单例级：模块内缓存）"""
    global _vector_store
    if _vector_store is None:
        Path(config.kb_vector_dir).mkdir(parents=True, exist_ok=True)
        _vector_store = Chroma(
            collection_name=_COLLECTION,
            embedding_function=_get_embeddings(),
            persist_directory=config.kb_vector_dir,
        )
    return _vector_store


def _invalidate_corpus() -> None:
    global _corpus_cache
    _corpus_cache = None


# ---------------------------------------------------------------------------
# 中文感知切分器
# ---------------------------------------------------------------------------
_CHINESE_SEPARATORS = ["\n\n", "\n", "。", "；", "！", "？", "，", " ", ""]


def _make_splitter(
    chunk_size: int | None = None, chunk_overlap: int | None = None
) -> RecursiveCharacterTextSplitter:
    """中文优先的分隔符顺序：段落 > 句子 > 逗号，避免从句子中间切开"""
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or config.kb_chunk_size,
        chunk_overlap=config.kb_chunk_overlap if chunk_overlap is None else chunk_overlap,
        separators=_CHINESE_SEPARATORS,
    )


def _split_long_block(content: str, section: str, page: int | None = None) -> list[dict[str, Any]]:
    """超长块二次切分，保留章节/页码元数据"""
    if len(content) <= config.kb_chunk_size * 1.2:
        return [{"content": content, "section": section, "page": page}]
    out = []
    for piece in _make_splitter().split_text(content):
        out.append({"content": piece, "section": section, "page": page})
    return out


# ---------------------------------------------------------------------------
# Markdown 结构解析（标题树切块）
# ---------------------------------------------------------------------------
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _parse_markdown_blocks(content: str) -> list[dict[str, Any]]:
    """按标题层级把文档切成语义块，每个块带章节路径（如：手册 > 电机过热排查）"""
    blocks: list[dict[str, Any]] = []
    stack: list[str] = []
    cur: list[str] = []
    cur_section = ""

    def flush() -> None:
        nonlocal cur
        if cur:
            text = "\n".join(cur).strip()
            if text:
                blocks.append({"content": text, "section": cur_section})
            cur = []

    for line in content.splitlines():
        m = _HEADING_RE.match(line.rstrip())
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            stack = stack[: level - 1] + [title]
            cur_section = " > ".join(stack)
            cur.append(line.strip())
        else:
            cur.append(line.rstrip())
    flush()

    out: list[dict[str, Any]] = []
    for b in blocks:
        out.extend(_split_long_block(b["content"], b["section"]))
    return out


# ---------------------------------------------------------------------------
# 各格式解析器
# ---------------------------------------------------------------------------
def _table_to_markdown(rows: list[list[str]]) -> str:
    """表格转 Markdown，保结构（表格切碎会毁语义，尽量整块保留）"""
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * cols) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _parse_text_blocks(content: str) -> list[dict[str, Any]]:
    """纯文本：按段落聚合，中文感知切分"""
    paragraphs = [p.strip() for p in content.splitlines() if p.strip()]
    if not paragraphs:
        return []
    blocks: list[dict[str, Any]] = []
    buf = ""
    for p in paragraphs:
        if buf and len(buf) + len(p) > config.kb_chunk_size:
            blocks.extend(_split_long_block(buf, "正文"))
            buf = ""
        buf = f"{buf}\n{p}" if buf else p
    if buf:
        blocks.extend(_split_long_block(buf, "正文"))
    return blocks


def _parse_html_blocks(content: str) -> list[dict[str, Any]]:
    """HTML：去掉标签/导航，还原标题结构后复用 Markdown 解析"""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "html.parser")
    lines: list[str] = []
    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        if tag.name.startswith("h"):
            lines.append("#" * int(tag.name[1]) + " " + text)
        else:
            lines.append(text)
    return _parse_markdown_blocks("\n".join(lines))


def _parse_pdf_blocks(data: bytes) -> tuple[list[dict[str, Any]], int]:
    """PDF：PyMuPDF 提正文（带页码）+ pdfplumber 提表格，保留版面结构"""
    import fitz
    import pdfplumber

    blocks: list[dict[str, Any]] = []
    total_chars = 0
    try:
        pdf = fitz.open(stream=data, filetype="pdf")
        for pno, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            total_chars += len(text)
            if text:
                blocks.extend(_split_long_block(text, f"第{pno}页", pno))
        pdf.close()
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"[知识库] PyMuPDF 解析失败，尝试 pdfplumber: {exc}")

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for pno, page in enumerate(pdf.pages, start=1):
                for tbl in page.extract_tables() or []:
                    md = _table_to_markdown([[str(c) if c else "" for c in row] for row in tbl])
                    if md:
                        blocks.append({"content": md, "section": f"第{pno}页 表格", "page": pno})
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"[知识库] PDF 表格提取失败: {exc}")
    return blocks, total_chars


def _parse_docx_blocks(data: bytes) -> list[dict[str, Any]]:
    """DOCX：按段落样式识别标题层级，表格整体保留"""
    import docx

    doc = docx.Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        m = re.match(r"(heading|标题)\s*(\d)", style)
        if m:
            lines.append("#" * min(int(m.group(2)), 6) + " " + text)
        else:
            lines.append(text)
    for table in doc.tables:
        md = _table_to_markdown([[cell.text.strip() for cell in row.cells] for row in table.rows])
        if md:
            lines.append(md)
    return _parse_markdown_blocks("\n".join(lines))


def _parse_json_blocks(data: bytes | str) -> list[dict[str, Any]]:
    """JSON：按键路径切块（如 root > device > threshold），保留层级可溯源"""
    if isinstance(data, bytes):
        data = data.decode("utf-8", errors="replace")
    obj = json.loads(data)
    blocks: list[dict[str, Any]] = []

    def walk(node: Any, path: str) -> None:
        if isinstance(node, dict):
            for k, v in node.items():
                walk(v, f"{path} > {k}" if path else str(k))
        elif isinstance(node, list):
            if len(node) <= 8:
                blocks.extend(_split_long_block(f"{path}: {json.dumps(node, ensure_ascii=False)}", path))
            else:
                for i, item in enumerate(node):
                    walk(item, f"{path}[{i}]")
        elif node is None or isinstance(node, (str, int, float, bool)):
            text = str(node)
            if len(text) > config.kb_chunk_size:
                blocks.extend(_split_long_block(text, path))
            else:
                blocks.append({"content": f"{path}: {text}", "section": path})

    walk(obj, "root")
    return blocks


def _parse_csv_blocks(data: bytes | str) -> list[dict[str, Any]]:
    """CSV：表头 + 行分组，整块保留列语义"""
    if isinstance(data, bytes):
        data = data.decode("utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(data)))
    if not rows:
        return []
    header = rows[0]
    blocks: list[dict[str, Any]] = []
    group_size = 20
    for start in range(1, len(rows), group_size):
        lines = [" | ".join(header)]
        for row in rows[start : start + group_size]:
            lines.append(" | ".join(str(c) for c in row))
        blocks.append({"content": "\n".join(lines), "section": "CSV 数据"})
    return blocks


def _parse_xlsx_blocks(data: bytes) -> list[dict[str, Any]]:
    """XLSX：按工作表转 Markdown 表格，按行组分块"""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    blocks: list[dict[str, Any]] = []
    try:
        for ws in wb.worksheets:
            rows = [["" if v is None else str(v) for v in row] for row in ws.iter_rows(values_only=True)]
            if not rows:
                continue
            group_size = 20
            for start in range(0, len(rows), group_size):
                md = _table_to_markdown(rows[start : start + group_size])
                if md:
                    blocks.append({"content": md, "section": f"工作表 {ws.title}"})
    finally:
        wb.close()
    return blocks


# ---------------------------------------------------------------------------
# 入库
# ---------------------------------------------------------------------------
SUFFIX_HANDLERS: dict[str, tuple[str, Callable]] = {
    ".md": ("MD", _parse_markdown_blocks),
    ".markdown": ("MD", _parse_markdown_blocks),
    ".txt": ("TXT", _parse_text_blocks),
    ".html": ("HTML", _parse_html_blocks),
    ".htm": ("HTML", _parse_html_blocks),
    ".pdf": ("PDF", _parse_pdf_blocks),
    ".docx": ("DOCX", _parse_docx_blocks),
    ".json": ("JSON", _parse_json_blocks),
    ".csv": ("CSV", _parse_csv_blocks),
    ".xlsx": ("XLSX", _parse_xlsx_blocks),
}


def _ingest(filename: str, doc_type: str, blocks: list[dict[str, Any]], created_by: str) -> dict[str, Any]:
    """公共入库逻辑：写向量库 + 记录 KbDocument"""
    if not blocks:
        return {"chunks": 0}
    docs = []
    for i, b in enumerate(blocks):
        meta: dict[str, Any] = {"source": filename, "chunk_index": i, "doc_type": doc_type}
        if b.get("section"):
            meta["section"] = b["section"]
        if b.get("page") is not None:
            meta["page"] = b["page"]
        docs.append(Document(page_content=b["content"], metadata=meta))

    store = _get_vector_store()
    store.add_documents(docs)
    _invalidate_corpus()
    logger.info(f"[知识库] 已入库 {filename}: {len(blocks)} 块 ({doc_type})")

    db = SessionLocal()
    try:
        record = KbDocument(
            filename=filename,
            doc_type=doc_type,
            chunk_count=len(blocks),
            status="READY",
            created_by=created_by,
        )
        db.add(record)
        db.commit()
        db.refresh(record)
        doc_id = record.id
    finally:
        db.close()
    return {"doc_id": doc_id, "chunks": len(blocks)}


def ingest_text(filename: str, content: str, doc_type: str = "MD", created_by: str | None = "admin") -> dict[str, Any]:
    """文本入库（兼容旧调用：seed_kb / seed_demo）"""
    doc_type = (doc_type or "TXT").upper()
    if doc_type == "MD":
        blocks = _parse_markdown_blocks(content)
    elif doc_type == "HTML":
        blocks = _parse_html_blocks(content)
    else:
        blocks = _parse_text_blocks(content)
    return _ingest(filename, doc_type, blocks, created_by or "admin")


def ingest_file(filename: str, content_bytes: bytes, created_by: str = "admin") -> dict[str, Any]:
    """按文件扩展名解析并入库（支持 md/txt/html/pdf/docx/json/csv/xlsx）"""
    suffix = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix not in SUFFIX_HANDLERS:
        raise ValueError(f"不支持的文件类型: {suffix or '(无扩展名)'}（支持 md/txt/html/pdf/docx/json/csv/xlsx）")
    doc_type, handler = SUFFIX_HANDLERS[suffix]

    if doc_type in ("MD", "TXT", "HTML"):
        content = content_bytes.decode("utf-8", errors="replace")
        blocks = handler(content)
    elif doc_type == "PDF":
        blocks, total_chars = handler(content_bytes)  # type: ignore[misc]
        result = _ingest(filename, doc_type, blocks, created_by)
        if total_chars < 100:
            result["warning"] = "该 PDF 提取文本过少，疑似扫描件，建议先 OCR 再上传"
            logger.warning(f"[知识库] {filename} 提取文本过少（{total_chars} 字符），疑似扫描件")
        return result
    else:
        blocks = handler(content_bytes)  # type: ignore[operator]
    return _ingest(filename, doc_type, blocks, created_by)


# ---------------------------------------------------------------------------
# 混合检索：向量（语义） + BM25（关键词）→ RRF 融合
# ---------------------------------------------------------------------------
_LATIN_RE = re.compile(r"[a-z0-9]+")
_CJK_RE = re.compile(r"[\u4e00-\u9fff]")


def _tokenize(text: str) -> list[str]:
    """轻量分词：英文/数字词 + 中文二元组（无 jieba 依赖下的近似方案）"""
    text = text.lower()
    toks = _LATIN_RE.findall(text)
    cjk = _CJK_RE.findall(text)
    toks.extend("".join(cjk[i : i + 2]) for i in range(len(cjk) - 1))
    return toks


def _load_corpus() -> list[dict[str, Any]]:
    """从 Chroma 拉取全部片段作为 BM25 语料（小知识库规模下内存足够）"""
    global _corpus_cache
    if _corpus_cache is None:
        try:
            got = _get_vector_store().get(include=["documents", "metadatas"])
            ids = got.get("ids") or []
            docs = got.get("documents") or []
            metas = got.get("metadatas") or []
            _corpus_cache = [
                {"content": str(docs[i]), "meta": metas[i] or {}} for i in range(len(ids))
            ]
        except Exception as exc:  # noqa: BLE001
            logger.warning(f"[知识库] 全文语料加载失败，回退纯向量检索: {exc}")
            _corpus_cache = []
    return _corpus_cache


def _bm25_scores(query: str, corpus: list[dict[str, Any]], k1: float = 1.5, b: float = 0.75) -> list[float]:
    """BM25 关键词打分（无 rank_bm25 依赖的自实现）"""
    n = len(corpus)
    if n == 0:
        return []
    lens = [len(c["content"]) for c in corpus]
    avgdl = sum(lens) / n
    q_terms = set(_tokenize(query))
    if not q_terms:
        return [0.0] * n
    df: Counter = Counter()
    doc_tf: list[Counter] = []
    for c in corpus:
        toks = _tokenize(c["content"])
        doc_tf.append(Counter(toks))
        for t in set(toks):
            df[t] += 1
    scores = []
    for i, tf in enumerate(doc_tf):
        s = 0.0
        dl = lens[i]
        for t in q_terms:
            f = tf.get(t, 0)
            if f == 0:
                continue
            idf = math.log(1 + (n - df[t] + 0.5) / (df[t] + 0.5))
            s += idf * (f * (k1 + 1)) / (f + k1 * (1 - b + b * dl / avgdl))
        scores.append(s)
    return scores


def _format_result(content: str, meta: dict[str, Any], score: float) -> dict[str, Any]:
    return {
        "content": content,
        "source": meta.get("source", "未知来源"),
        "doc_type": meta.get("doc_type", ""),
        "section": meta.get("section", ""),
        "page": meta.get("page"),
        "score": round(score, 4),
    }


def search(query: str, k: int | None = None) -> list[dict[str, Any]]:
    """混合检索：向量 top-N + BM25 top-N，RRF 融合后取 top-k"""
    k = k or config.kb_top_k
    store = _get_vector_store()
    vector_results = store.similarity_search_with_relevance_scores(query, k=max(k * 3, 20))
    corpus = _load_corpus()
    if not corpus:
        # 语料不可用时回退纯向量
        return [
            _format_result(doc.page_content, doc.metadata, score)
            for doc, score in vector_results[:k]
        ]

    content_idx = {c["content"]: i for i, c in enumerate(corpus)}
    bm = _bm25_scores(query, corpus)
    bm_rank = sorted(range(len(corpus)), key=lambda i: bm[i], reverse=True)[: max(k * 3, 20)]

    rrf: dict[int, float] = {}
    for rank, (doc, _score) in enumerate(vector_results):
        idx = content_idx.get(doc.page_content)
        if idx is not None:
            rrf[idx] = rrf.get(idx, 0.0) + 1.0 / (60 + rank)
    for rank, idx in enumerate(bm_rank):
        rrf[idx] = rrf.get(idx, 0.0) + 1.0 / (60 + rank)

    top = sorted(rrf.items(), key=lambda kv: kv[1], reverse=True)[:k]
    items = []
    for idx, s in top:
        meta = corpus[idx]["meta"]
        # RRF 分数归一化到 0~1（两路满分 = 2/60）
        items.append(_format_result(corpus[idx]["content"], meta, min(s * 30.0, 1.0)))
    return items


def search_as_context(query: str, k: int | None = None) -> str:
    """检索并格式化为 LLM 友好的上下文（带来源/章节/页码）"""
    items = search(query, k)
    if not items:
        return "没有找到相关资料。"
    parts = []
    for i, item in enumerate(items, 1):
        header = f"【资料 {i} | 来源: {item['source']}"
        if item.get("section"):
            header += f" | {item['section']}"
        if item.get("page") is not None:
            header += f" | 第{item['page']}页"
        parts.append(f"{header}】\n{item['content']}")
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# 文档管理
# ---------------------------------------------------------------------------
def list_documents(db: Session) -> list[dict[str, Any]]:
    rows = db.scalars(select(KbDocument).order_by(KbDocument.created_at.desc())).all()
    return [
        {
            "id": row.id,
            "filename": row.filename,
            "doc_type": row.doc_type,
            "chunk_count": row.chunk_count,
            "status": row.status,
            "created_by": row.created_by,
            "created_at": row.created_at.isoformat() if row.created_at else None,
        }
        for row in rows
    ]


def delete_document(db: Session, doc_id: int) -> bool:
    """删除文档：清 Chroma 中该来源向量 + 删记录"""
    record = db.get(KbDocument, doc_id)
    if record is None:
        return False
    store = _get_vector_store()
    try:
        store.delete(where={"source": record.filename})
    except Exception as exc:  # noqa: BLE001
        logger.warning(f"[知识库] 删除向量失败（可能已不存在）: {exc}")
    db.delete(record)
    db.commit()
    _invalidate_corpus()
    return True


def seed_demo(db: Session) -> dict[str, Any]:
    """灌入预置演示手册（已存在则跳过）"""
    doc_path = Path(__file__).resolve().parents[2] / "kb_docs" / "设备维保手册.md"
    if not doc_path.exists():
        return {"seeded": False, "reason": "演示手册文件不存在"}
    exists = db.scalar(select(KbDocument.id).where(KbDocument.filename == doc_path.name).limit(1))
    if exists is not None:
        return {"seeded": False, "reason": "演示手册已存在"}
    result = ingest_text(doc_path.name, doc_path.read_text(encoding="utf-8"), doc_type="MD", created_by="seed")
    return {"seeded": True, **result}
